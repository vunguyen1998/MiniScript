from docx import Document
import os
import psutil
import pygetwindow as gw

document = Document()

def create_beta_folder(truyen_path):
    """Tạo thư mục BETA nếu chưa tồn tại trong thư mục truyen"""
    
    beta_folder_path = os.path.join(truyen_path, "BETA")
    if not os.path.exists(beta_folder_path):
        os.makedirs(beta_folder_path)
        print("Thư mục BETA đã được tạo.")
    else:
        print("Thư mục BETA đã tồn tại.")

def read_file(truyen, header, contents, CV_chapt_from, CV_chapt_to):
    for chuong in range(CV_chapt_from, int(CV_chapt_to + 1)):
        if not os.path.exists(f"{truyen}/Chương {chuong}.docx"):
            print(f"Lỗi: Tệp '{truyen}/Chương {chuong}.docx' không tồn tại.")
            messagebox.showerror("Lỗi", f"Tệp '{truyen}/Chương {chuong}.docx' \nkhông tồn tại.")
            break
        else:
            document = Document(f"{truyen}/Chương {chuong}.docx") #read file
            lines = document.paragraphs
            line_num = 1
            for line in lines:
                if line_num == 1:
                    header.append(line.text.strip('. ').split(': ')[1].title())
                else:
                    if line.text != '':
                        contents.append(line.text.strip())
                line_num += 1
        
def read_beta_file(truyen, BETA_chapt, contents):
    try:
        document = Document(f"{truyen}/BETA/Chương {BETA_chapt}.docx") #read file
        lines = document.paragraphs
        line_num = 1
        for line in lines:
            if line_num != 1:
                if line.text != '':
                    contents.append(line.text.strip())
            line_num += 1
    except:
        contents.append('')
        
def open_or_focus_folder(folder_path):
    """Kiểm tra xem Windows Explorer có đang mở thư mục đó không

    Args:
        folder_path (str): Đường dẫn đến thư mục BETA
    """
    # Kiểm tra xem Windows Explorer có đang mở thư mục đó không
    folder_opened = False
    for proc in psutil.process_iter(['pid', 'name', 'cmdline']):
        if proc.info['name'] == 'explorer.exe' and proc.info['cmdline']:
            # Kiểm tra thư mục có trong lệnh cmdline của explorer.exe không
            if any(folder_path in cmd for cmd in proc.info['cmdline']):
                folder_opened = True
                break
    
    if folder_opened:
        # Nếu thư mục đã mở, tìm cửa sổ và đưa nó lên đầu
        for window in gw.getWindowsWithTitle(folder_path):
            window.activate()  # Đưa cửa sổ này lên đầu
            return
    else:
        # Nếu thư mục chưa mở, mở nó bằng os.startfile
        os.startfile(folder_path)
                           

def split_text_file(truyen = '', target_word_count = 0, BETA_chapt = 0, CV_chapt_from = 0, CV_chapt_to = 0):
    """Tách chương theo số chữ đặt trước

    Args:
        truyen (string): Đường dẫn đến folder chứa chương truyện đã Edit
        target_word_count (number): Số chữ tối thiểu mỗi chương, = 0 là không tách chương
        BETA_chapt (number): Chương BETA mới nhất
        CV_chapt_from (number): Chương đã Edit bắt đầu từ
        CV_chapt_to (number): Chương đã Edit kết thúc ở
    """
    
    create_beta_folder(truyen)
    
    if target_word_count != 0:
        part = 1
        current_words = 0
        current_content = []
        header = []
        contents = []
        
        read_beta_file(truyen, BETA_chapt, contents)
        read_file(truyen, header, contents, CV_chapt_from, CV_chapt_to)
        
        # Ước tính số chương BETA
        words = 0
        for line in contents: 
            words += len(line.split())
        so_chuong_beta = words//target_word_count
        print(so_chuong_beta)
        progress_bar["maximum"] = so_chuong_beta
        
        
        header_after = []
        for i in range(len(header)):
            if len(header) != so_chuong_beta:
                header_after.append(header[i] + ' (1)')
                header_after.append(header[i] + ' (2)')
                if so_chuong_beta//len(header) >= 2:
                    header_after.append(header[i] + ' (3)')
                if so_chuong_beta//len(header) >= 3:
                    header_after.append(header[i] + ' (4)')
                if so_chuong_beta//len(header) >= 3:
                    header_after.append(header[i] + ' (5)')
                if so_chuong_beta//len(header) >= 3:
                    header_after.append(header[i] + ' (6)')
                so_chuong_beta = so_chuong_beta - 1
            else:
                header_after.append(header[i])
        header = header_after
        print(header)

        for line in contents:
            word_count = len(line.split())
            if current_words + word_count <= target_word_count:
                current_content.append(line)
                current_words += word_count
            else:
                # Lưu file nếu số từ đạt yêu cầu
                document = Document()
                document.add_paragraph(f"Chương {BETA_chapt}: " + header[int(part-1)])
                for i in current_content:
                    document.add_paragraph(i)
                document.save(f"{truyen}/BETA/Chương {BETA_chapt}.docx")
                
                # Status bar
                progress_bar["value"] = part
                progress_bar.update()
                
                # Tăng số thứ tự file và reset các biến đếm
                part += 1
                BETA_chapt += 1
                current_content = [line]
                current_words = word_count
            
        # Lưu file cuối cùng nếu còn nội dung
        if current_content:
            document = Document()
            document.add_paragraph([f"Chương {BETA_chapt}: "])
            for i in current_content:
                document.add_paragraph(i)
            document.save(f"{truyen}/BETA/Chương {BETA_chapt}.docx")
            

    else:
        part = 1
        progress_bar["maximum"] = CV_chapt_to - CV_chapt_from + 1
        for chuong in range(CV_chapt_from, int(CV_chapt_to + 1)):
            if not os.path.exists(f"{truyen}/Chương {chuong}.docx"):
                print(f"Lỗi: Tệp '{truyen}/Chương {chuong}.docx' không tồn tại.")
                messagebox.showerror("Lỗi", f"Tệp '{truyen}/Chương {chuong}.docx' \nkhông tồn tại.")
                break
            else:
                document = Document(f"{truyen}/Chương {chuong}.docx") #read file
                lines = document.paragraphs
                line_num = 1
                document = Document()
                BETA_chapt += 1
                for line in lines:
                    if line_num == 1:
                        document.add_paragraph(f'Chương {BETA_chapt}: ' + line.text.strip('. ').split(': ')[1].title())
                    else:
                        if line.text != '':
                            document.add_paragraph(line.text.strip())
                    line_num += 1
                    
                document.save(f"{truyen}/BETA/Chương {BETA_chapt}.docx")
                # Status bar
                progress_bar["value"] = part
                progress_bar.update()
                
                part += 1
        
    # Hiện thông báo hoàn thành
    messagebox.showinfo("Hoàn thành", f"Quá trình xử lý hoàn tất!\n Chương được lưu ở {truyen}/BETA")
    open_or_focus_folder(f"{truyen}/BETA")

# GUI #
import os
import glob
import re
from tkinter import Tk, filedialog, StringVar, IntVar, messagebox
from tkinter import ttk

def get_latest_chapter(folder_path, prefix="Chương"):
    """Tìm chương mới nhất từ tên file trong folder"""
    files = glob.glob(os.path.join(folder_path, f"{prefix}*.docx"))
    chapter_numbers = []
    for file in files:
        match = re.search(rf"{prefix}\s*(\d+)", os.path.basename(file))
        if match:
            chapter_numbers.append(int(match.group(1)))
    return max(chapter_numbers) if chapter_numbers else None

def browse_folder():
    """Chọn thư mục truyện"""
    folder_path = filedialog.askdirectory()
    truyen_folder.set(folder_path)

    # Lấy số chương BETA mới nhất
    beta_folder = os.path.join(folder_path, "BETA")
    latest_beta_chapter = get_latest_chapter(beta_folder) if os.path.exists(beta_folder) else None
    if latest_beta_chapter:
        beta_chapter_entry.delete(0, 'end')
        beta_chapter_entry.insert(0, str(latest_beta_chapter))
    
    # Lấy số chương CV mới nhất
    latest_cv_chapter = get_latest_chapter(folder_path)
    if latest_cv_chapter:
        cv_chapt_to_entry.delete(0, 'end')
        cv_chapt_to_entry.insert(0, str(latest_cv_chapter))

def toggle_target_word_count():
    """Hiện hoặc ẩn input nhập số chữ mỗi chương dựa trên lựa chọn"""
    if target_word_var.get() == 1:
        word_count_entry.config(state="normal")
    else:
        word_count_entry.config(state="disabled")

def split_text_file_gui():
    """Hàm tách chương xử lý theo giao diện nhập"""
    truyen = truyen_folder.get()
    target_word_count = 0 if target_word_var.get() == 0 else int(word_count_entry.get())
    beta_chapter = int(beta_chapter_entry.get())
    cv_chapt_from = int(cv_chapt_from_entry.get())
    cv_chapt_to = int(cv_chapt_to_entry.get())

    print(f"Truyện: {truyen}")
    print(f"Tách chương: {target_word_count}")
    print(f"BETA chương: {beta_chapter}")
    print(f"CV chương từ: {cv_chapt_from} đến {cv_chapt_to}")
    
    # Khởi động thanh tiến trình
    progress_bar["value"] = 0
    progress_bar.update()
    split_text_file(truyen, target_word_count, beta_chapter, cv_chapt_from, cv_chapt_to)


# Tạo giao diện
root = Tk()
root.title("Tách Chương Truyện")
root.geometry("500x550")
root.configure(bg="#F0F0F0")

# function to validate mark entry
def only_numbers(char):
    return char.isdigit()
validation = root.register(only_numbers)

style = ttk.Style()
style.configure("TLabel", font=("Arial", 10), background="#F0F0F0")
style.configure("TButton", font=("Arial", 10), padding=5)
style.configure("TEntry", padding=5)
style.configure("TRadiobutton", background="#F0F0F0")

# Chọn thư mục truyện
frame_folder = ttk.Frame(root, padding=10)
frame_folder.pack(fill="x", padx=20, pady=5)
ttk.Label(frame_folder, text="Chọn thư mục truyện:").grid(row=0, column=0, sticky="w")
truyen_folder = StringVar()
folder_label = ttk.Label(frame_folder, textvariable=truyen_folder, font=("Arial", 10, "italic"))
folder_label.grid(row=1, column=0, sticky="w", pady=(5, 0))
ttk.Button(frame_folder, text="Browse", command=browse_folder).grid(row=0, column=1, padx=10, sticky="e")

# Tùy chọn tách chương
frame_options = ttk.Frame(root, padding=10)
frame_options.pack(fill="x", padx=20, pady=5)
ttk.Label(frame_options, text="Tách chương:").pack(anchor="w")
target_word_var = IntVar(value=0)
ttk.Radiobutton(frame_options, text="Không tách chương", variable=target_word_var, value=0, command=toggle_target_word_count).pack(anchor="w")
ttk.Radiobutton(frame_options, text="Tách chương", variable=target_word_var, value=1, command=toggle_target_word_count).pack(anchor="w")

# Nhập số từ mỗi chương
frame_word_count = ttk.Frame(root, padding=10)
frame_word_count.pack(fill="x", padx=20, pady=5)
ttk.Label(frame_word_count, text="Số chữ mỗi chương:").grid(row=0, column=0, sticky="w")
word_count_entry = ttk.Entry(frame_word_count, validate="key", validatecommand=(validation, '%S'))
word_count_entry.insert(0, "0")
word_count_entry.config(state="disabled")
word_count_entry.grid(row=0, column=1, padx=5)

# Nhập chương BETA
frame_beta = ttk.Frame(root, padding=10)
frame_beta.pack(fill="x", padx=20, pady=5)
ttk.Label(frame_beta, text="Chương BETA:").grid(row=0, column=0, sticky="w")
beta_chapter_entry = ttk.Entry(frame_beta, validate="key", validatecommand=(validation, '%S'))
beta_chapter_entry.grid(row=0, column=1, padx=5)

# Nhập chương CV từ
frame_cv_from = ttk.Frame(root, padding=10)
frame_cv_from.pack(fill="x", padx=20, pady=5)
ttk.Label(frame_cv_from, text="Chương CV từ:").grid(row=0, column=0, sticky="w")
cv_chapt_from_entry = ttk.Entry(frame_cv_from, validate="key", validatecommand=(validation, '%S'))
cv_chapt_from_entry.grid(row=0, column=1, padx=5)

# Nhập chương CV đến
frame_cv_to = ttk.Frame(root, padding=10)
frame_cv_to.pack(fill="x", padx=20, pady=5)
ttk.Label(frame_cv_to, text="Chương CV đến:").grid(row=0, column=0, sticky="w")
cv_chapt_to_entry = ttk.Entry(frame_cv_to, validate="key", validatecommand=(validation, '%S'))
cv_chapt_to_entry.grid(row=0, column=1, padx=5)

# Nút Thực hiện và Thanh tiến trình
frame_button = ttk.Frame(root, padding=10)
frame_button.pack(pady=10)
ttk.Button(frame_button, text="Thực hiện", command=split_text_file_gui).pack()
progress_bar = ttk.Progressbar(frame_button, orient="horizontal", length=300, mode="determinate")
progress_bar.pack(pady=10)

root.mainloop()